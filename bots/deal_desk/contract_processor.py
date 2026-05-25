#!/usr/bin/env python3
"""
Contract Template Processor for Realtor Agent
Processes Word document templates and communicates with the Realtor Agent application.
"""

import os
import sys
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
import requests
from dataclasses import dataclass

# Add the project root to Python path
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not available. Install with: pip install python-docx")
    Document = None
    DOCX_AVAILABLE = False

from realtor_agent.core.config import load_config
from realtor_agent.utils.logging import setup_logging


@dataclass
class ContractData:
    """Data structure for contract generation."""
    property_id: str
    buyer_name: str
    buyer_address: str
    seller_name: str
    seller_address: str
    property_address: str
    purchase_price: float
    earnest_money: float
    closing_date: str
    contingencies: List[str]
    special_terms: str = ""
    realtor_agent: str = "Realtor Agent LLC"
    agent_address: str = "123 Real Estate Ave, Suite 100"
    agent_phone: str = "(555) 123-4567"
    agent_email: str = "deals@realtoragent.com"


class ContractTemplateProcessor:
    """Processes contract templates and communicates with Realtor Agent application."""

    def __init__(self, config_path: str = "realtor_agent/agent_config.yml",
                 api_url: str = None, api_key: str = None):
        """
        Initialize the contract processor.

        Args:
            config_path: Path to Realtor Agent configuration
            api_url: API URL for the Realtor Agent application
            api_key: API key for authentication
        """
        self.config_path = config_path
        self.api_url = api_url or os.getenv('REALTOR_AGENT_API_URL', 'http://localhost:5000')
        self.api_key = api_key or os.getenv('REALTOR_AGENT_API_KEY', 'test_key')

        # Setup logging
        setup_logging()
        self.logger = logging.getLogger(__name__)

        # Load configuration
        try:
            self.config = load_config(Path(config_path))
        except Exception as e:
            self.logger.warning(f"Could not load config: {e}")
            self.config = None

        # Template directory
        self.template_dir = Path(__file__).parent / "contracts"
        self.output_dir = Path("generated_contracts")
        self.output_dir.mkdir(exist_ok=True)

        self.logger.info("ContractTemplateProcessor initialized")

    def get_deal_data_from_api(self, deal_id: str) -> Optional[Dict[str, Any]]:
        """
        Retrieve deal data from Realtor Agent API.

        Args:
            deal_id: Deal identifier

        Returns:
            Deal data dictionary or None if not found
        """
        try:
            url = f"{self.api_url}/api/deals/{deal_id}"
            headers = {"X-API-Key": self.api_key}

            response = requests.get(url, headers=headers, timeout=30)

            if response.status_code == 200:
                deal_data = response.json()
                self.logger.info(f"Retrieved deal data for {deal_id}")
                return deal_data
            else:
                self.logger.error(f"Failed to get deal data: {response.status_code} - {response.text}")
                return None

        except requests.exceptions.RequestException as e:
            self.logger.error(f"API request failed: {e}")
            return None

    def create_contract_data_from_deal(self, deal_data: Dict[str, Any]) -> ContractData:
        """
        Convert deal data to contract data structure.

        Args:
            deal_data: Deal data from API

        Returns:
            ContractData object
        """
        # Extract property information
        property_info = deal_data.get('property', {})
        property_address = property_info.get('address', 'Property Address Not Available')

        # Extract buyer information
        buyer_info = deal_data.get('buyer', {})
        buyer_name = buyer_info.get('name', 'Buyer Name Not Available')
        buyer_address = buyer_info.get('address', 'Buyer Address Not Available')

        # Extract seller information
        seller_info = deal_data.get('seller', {})
        seller_name = seller_info.get('name', 'Seller Name Not Available')
        seller_address = seller_info.get('address', 'Seller Address Not Available')

        # Extract financial information
        financial_info = deal_data.get('financial', {})
        purchase_price = financial_info.get('purchase_price', 0)
        earnest_money = financial_info.get('earnest_money', purchase_price * 0.01)  # Default 1%

        # Extract terms
        terms = deal_data.get('terms', {})
        closing_date = terms.get('closing_date', '30 days from acceptance')
        contingencies = terms.get('contingencies', ['inspection', 'appraisal'])
        special_terms = terms.get('special_terms', '')

        return ContractData(
            property_id=deal_data.get('id', 'unknown'),
            buyer_name=buyer_name,
            buyer_address=buyer_address,
            seller_name=seller_name,
            seller_address=seller_address,
            property_address=property_address,
            purchase_price=purchase_price,
            earnest_money=earnest_money,
            closing_date=closing_date,
            contingencies=contingencies,
            special_terms=special_terms
        )

    def process_template(self, template_path: Path, contract_data: ContractData,
                        output_path: Path) -> bool:
        """
        Process a Word document template with contract data.

        Args:
            template_path: Path to template document
            contract_data: Contract data to fill in
            output_path: Path for output document

        Returns:
            True if successful, False otherwise
        """
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available. Cannot process Word templates.")
            return False

        try:
            # Load template
            doc = Document(template_path)

            # Create replacement mapping
            replacements = self._create_replacement_mapping(contract_data)

            # Process all paragraphs
            for paragraph in doc.paragraphs:
                self._replace_text_in_paragraph(paragraph, replacements)

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_text_in_paragraph(paragraph, replacements)

            # Save processed document
            doc.save(output_path)
            self.logger.info(f"Contract generated: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error processing template: {e}")
            return False

    def _create_replacement_mapping(self, contract_data: ContractData) -> Dict[str, str]:
        """Create mapping of template variables to actual values."""
        return {
            "{{BUYER_NAME}}": contract_data.buyer_name,
            "{{BUYER_ADDRESS}}": contract_data.buyer_address,
            "{{SELLER_NAME}}": contract_data.seller_name,
            "{{SELLER_ADDRESS}}": contract_data.seller_address,
            "{{PROPERTY_ADDRESS}}": contract_data.property_address,
            "{{PURCHASE_PRICE}}": f"${contract_data.purchase_price:,.2f}",
            "{{EARNEST_MONEY}}": f"${contract_data.earnest_money:,.2f}",
            "{{CLOSING_DATE}}": contract_data.closing_date,
            "{{CONTINGENCIES}}": ", ".join(contract_data.contingencies),
            "{{SPECIAL_TERMS}}": contract_data.special_terms or "None",
            "{{AGENT_NAME}}": contract_data.realtor_agent,
            "{{AGENT_ADDRESS}}": contract_data.agent_address,
            "{{AGENT_PHONE}}": contract_data.agent_phone,
            "{{AGENT_EMAIL}}": contract_data.agent_email,
            "{{CURRENT_DATE}}": datetime.now().strftime("%B %d, %Y"),
            "{{CONTRACT_DATE}}": datetime.now().strftime("%B %d, %Y")
        }

    def _replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """Replace template variables in a paragraph."""
        if not hasattr(paragraph, 'text'):
            return

        original_text = paragraph.text
        for placeholder, value in replacements.items():
            if placeholder in original_text:
                original_text = original_text.replace(placeholder, str(value))

        paragraph.text = original_text

    def generate_cash_offer_contract(self, deal_id: str, output_filename: str = None) -> Optional[Path]:
        """
        Generate a cash offer contract for a specific deal.

        Args:
            deal_id: Deal identifier
            output_filename: Optional output filename

        Returns:
            Path to generated contract or None if failed
        """
        self.logger.info(f"Generating cash offer contract for deal {deal_id}")

        # Get deal data from API
        deal_data = self.get_deal_data_from_api(deal_id)
        if not deal_data:
            self.logger.error(f"Could not retrieve deal data for {deal_id}")
            return None

        # Convert to contract data
        contract_data = self.create_contract_data_from_deal(deal_data)

        # Template path
        template_path = self.template_dir / "cash_offer_template.docx"
        if not template_path.exists():
            self.logger.error(f"Template not found: {template_path}")
            return None

        # Output path
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"cash_offer_{deal_id}_{timestamp}.docx"

        output_path = self.output_dir / output_filename

        # Process template
        if self.process_template(template_path, contract_data, output_path):
            self.logger.info(f"Cash offer contract generated: {output_path}")
            return output_path
        else:
            self.logger.error("Failed to generate cash offer contract")
            return None

    def generate_lease_option_contract(self, deal_id: str, output_filename: str = None) -> Optional[Path]:
        """
        Generate a lease option contract for a specific deal.

        Args:
            deal_id: Deal identifier
            output_filename: Optional output filename

        Returns:
            Path to generated contract or None if failed
        """
        self.logger.info(f"Generating lease option contract for deal {deal_id}")

        # Get deal data from API
        deal_data = self.get_deal_data_from_api(deal_id)
        if not deal_data:
            self.logger.error(f"Could not retrieve deal data for {deal_id}")
            return None

        # Convert to contract data
        contract_data = self.create_contract_data_from_deal(deal_data)

        # Template path
        template_path = self.template_dir / "lease_option_template.docx"
        if not template_path.exists():
            self.logger.error(f"Template not found: {template_path}")
            return None

        # Output path
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"lease_option_{deal_id}_{timestamp}.docx"

        output_path = self.output_dir / output_filename

        # Process template
        if self.process_template(template_path, contract_data, output_path):
            self.logger.info(f"Lease option contract generated: {output_path}")
            return output_path
        else:
            self.logger.error("Failed to generate lease option contract")
            return None

    def generate_owner_finance_contract(self, deal_id: str, output_filename: str = None) -> Optional[Path]:
        """
        Generate an owner finance term sheet for a specific deal.

        Args:
            deal_id: Deal identifier
            output_filename: Optional output filename

        Returns:
            Path to generated contract or None if failed
        """
        self.logger.info(f"Generating owner finance contract for deal {deal_id}")

        # Get deal data from API
        deal_data = self.get_deal_data_from_api(deal_id)
        if not deal_data:
            self.logger.error(f"Could not retrieve deal data for {deal_id}")
            return None

        # Convert to contract data
        contract_data = self.create_contract_data_from_deal(deal_data)

        # Template path
        template_path = self.template_dir / "owner_finance_term_sheet.docx"
        if not template_path.exists():
            self.logger.error(f"Template not found: {template_path}")
            return None

        # Output path
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"owner_finance_{deal_id}_{timestamp}.docx"

        output_path = self.output_dir / output_filename

        # Process template
        if self.process_template(template_path, contract_data, output_path):
            self.logger.info(f"Owner finance contract generated: {output_path}")
            return output_path
        else:
            self.logger.error("Failed to generate owner finance contract")
            return None

    def generate_subject_to_contract(self, deal_id: str, output_filename: str = None) -> Optional[Path]:
        """
        Generate a subject-to contract for a specific deal.

        Args:
            deal_id: Deal identifier
            output_filename: Optional output filename

        Returns:
            Path to generated contract or None if failed
        """
        self.logger.info(f"Generating subject-to contract for deal {deal_id}")

        # Get deal data from API
        deal_data = self.get_deal_data_from_api(deal_id)
        if not deal_data:
            self.logger.error(f"Could not retrieve deal data for {deal_id}")
            return None

        # Convert to contract data
        contract_data = self.create_contract_data_from_deal(deal_data)

        # Template path
        template_path = self.template_dir / "subject_to_template.docx"
        if not template_path.exists():
            self.logger.error(f"Template not found: {template_path}")
            return None

        # Output path
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"subject_to_{deal_id}_{timestamp}.docx"

        output_path = self.output_dir / output_filename

        # Process template
        if self.process_template(template_path, contract_data, output_path):
            self.logger.info(f"Subject-to contract generated: {output_path}")
            return output_path
        else:
            self.logger.error("Failed to generate subject-to contract")
            return None

    def upload_contract_to_api(self, contract_path: Path, deal_id: str,
                              contract_type: str) -> bool:
        """
        Upload generated contract to Realtor Agent API.

        Args:
            contract_path: Path to generated contract
            deal_id: Deal identifier
            contract_type: Type of contract (cash_offer, lease_option, etc.)

        Returns:
            True if upload successful, False otherwise
        """
        try:
            url = f"{self.api_url}/api/contracts"
            headers = {"X-API-Key": self.api_key}

            # Prepare multipart form data
            with open(contract_path, 'rb') as f:
                files = {
                    'contract_file': (contract_path.name, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                }
                data = {
                    'deal_id': deal_id,
                    'contract_type': contract_type,
                    'generated_at': datetime.now().isoformat()
                }

                response = requests.post(url, files=files, data=data, headers=headers, timeout=60)

                if response.status_code in [200, 201]:
                    self.logger.info(f"Contract uploaded successfully for deal {deal_id}")
                    return True
                else:
                    self.logger.error(f"Failed to upload contract: {response.status_code} - {response.text}")
                    return False

        except Exception as e:
            self.logger.error(f"Error uploading contract: {e}")
            return False

    def list_available_templates(self) -> List[str]:
        """List all available contract templates."""
        templates = []
        if self.template_dir.exists():
            for template_file in self.template_dir.glob("*.docx"):
                templates.append(template_file.stem)
        return templates

    def validate_template_variables(self, template_path: Path) -> List[str]:
        """
        Validate and extract template variables from a document.

        Args:
            template_path: Path to template document

        Returns:
            List of template variables found
        """
        if not DOCX_AVAILABLE:
            return []

        variables = []
        try:
            doc = Document(template_path)

            # Check paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                # Find variables in {{VARIABLE}} format
                import re
                matches = re.findall(r'\{\{([^}]+)\}\}', text)
                variables.extend(matches)

            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            matches = re.findall(r'\{\{([^}]+)\}\}', text)
                            variables.extend(matches)

            # Remove duplicates
            variables = list(set(variables))

        except Exception as e:
            self.logger.error(f"Error validating template: {e}")

        return variables


def main():
    """Command-line interface for contract processing."""
    import argparse

    parser = argparse.ArgumentParser(description="Realtor Agent Contract Template Processor")
    parser.add_argument("--config", default="realtor_agent/agent_config.yml",
                       help="Path to Realtor Agent configuration")
    parser.add_argument("--api-url", help="Realtor Agent API URL")
    parser.add_argument("--api-key", help="API key for authentication")
    parser.add_argument("--deal-id", required=True, help="Deal ID to process")

    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    # Generate commands
    generate_parser = subparsers.add_parser("generate", help="Generate contracts")
    generate_subparsers = generate_parser.add_subparsers(dest="contract_type")

    # Cash offer
    cash_parser = generate_subparsers.add_parser("cash-offer", help="Generate cash offer contract")
    cash_parser.add_argument("--output", help="Output filename")

    # Lease option
    lease_parser = generate_subparsers.add_parser("lease-option", help="Generate lease option contract")
    lease_parser.add_argument("--output", help="Output filename")

    # Owner finance
    finance_parser = generate_subparsers.add_parser("owner-finance", help="Generate owner finance term sheet")
    finance_parser.add_argument("--output", help="Output filename")

    # Subject to
    subject_parser = generate_subparsers.add_parser("subject-to", help="Generate subject-to contract")
    subject_parser.add_argument("--output", help="Output filename")

    # Upload command
    upload_parser = subparsers.add_parser("upload", help="Upload contract to API")
    upload_parser.add_argument("--file", required=True, help="Contract file to upload")
    upload_parser.add_argument("--type", required=True, help="Contract type")

    # List templates
    subparsers.add_parser("list-templates", help="List available templates")

    # Validate template
    validate_parser = subparsers.add_parser("validate", help="Validate template variables")
    validate_parser.add_argument("--template", required=True, help="Template file to validate")

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return

    # Initialize processor
    processor = ContractTemplateProcessor(
        config_path=args.config,
        api_url=args.api_url,
        api_key=args.api_key
    )

    try:
        if args.command == "generate":
            if args.contract_type == "cash-offer":
                result = processor.generate_cash_offer_contract(args.deal_id, args.output)
                if result:
                    print(f"✅ Cash offer contract generated: {result}")
                    # Optionally upload
                    upload = input("Upload to API? (y/N): ").lower().strip()
                    if upload == 'y':
                        if processor.upload_contract_to_api(result, args.deal_id, "cash_offer"):
                            print("✅ Contract uploaded successfully")
                        else:
                            print("❌ Failed to upload contract")
                else:
                    print("❌ Failed to generate contract")
                    sys.exit(1)

            elif args.contract_type == "lease-option":
                result = processor.generate_lease_option_contract(args.deal_id, args.output)
                if result:
                    print(f"✅ Lease option contract generated: {result}")
                else:
                    print("❌ Failed to generate contract")
                    sys.exit(1)

            elif args.contract_type == "owner-finance":
                result = processor.generate_owner_finance_contract(args.deal_id, args.output)
                if result:
                    print(f"✅ Owner finance contract generated: {result}")
                else:
                    print("❌ Failed to generate contract")
                    sys.exit(1)

            elif args.contract_type == "subject-to":
                result = processor.generate_subject_to_contract(args.deal_id, args.output)
                if result:
                    print(f"✅ Subject-to contract generated: {result}")
                else:
                    print("❌ Failed to generate contract")
                    sys.exit(1)

        elif args.command == "upload":
            contract_path = Path(args.file)
            if not contract_path.exists():
                print(f"❌ Contract file not found: {args.file}")
                sys.exit(1)

            if processor.upload_contract_to_api(contract_path, args.deal_id, args.type):
                print("✅ Contract uploaded successfully")
            else:
                print("❌ Failed to upload contract")
                sys.exit(1)

        elif args.command == "list-templates":
            templates = processor.list_available_templates()
            if templates:
                print("Available templates:")
                for template in templates:
                    print(f"  - {template}")
            else:
                print("No templates found")

        elif args.command == "validate":
            template_path = Path(args.template)
            if not template_path.exists():
                print(f"❌ Template file not found: {args.template}")
                sys.exit(1)

            variables = processor.validate_template_variables(template_path)
            if variables:
                print(f"Template variables found in {args.template}:")
                for var in sorted(variables):
                    print(f"  - {{{{{var}}}}}")
            else:
                print(f"No template variables found in {args.template}")

    except Exception as e:
        print(f"❌ Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()